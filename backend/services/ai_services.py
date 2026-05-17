import io
import asyncio
import json
import logging
import re
import base64
import pypdfium2 as pdfium
from docx import Document
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from core.config import Config
from core.database import supabase
from services.code_lookup import (
    describe_diagnosis,
    describe_procedure,
    format_code_with_description,
)

logger = logging.getLogger(__name__)

# --- PROMPTS ---
# Pre-authorisation has no AI review step. The pre-auth pipeline OCRs the
# dropped-off documents and extracts structured fields only; every routed
# pre-auth goes to the insurer's manual review queue. (The claim scrubber and
# adjudication prompts are defined below.)


# --- LLM INITIALIZATION ---
def get_vision_llm():
    return ChatGoogleGenerativeAI(
        model=Config.OCR_MODEL, 
        google_api_key=Config.GEMINI_API_KEY, 
        temperature=0.1
    )

def get_llm(json_mode: bool = False):
    kwargs = {
        "api_key": Config.GROQ_API_KEY,
        "model_name": Config.LLM_MODEL,
        "temperature": 0.1,
    }
    if json_mode:
        # Groq honours OpenAI-style response_format. Guarantees the model
        # returns a JSON-parseable string.
        kwargs["model_kwargs"] = {"response_format": {"type": "json_object"}}
    return ChatGroq(**kwargs)

def get_embeddings():
    # Direct HTTP wrapper around Gemini's `embedContent` endpoint.
    # See services/embeddings.py for why we don't use GoogleGenerativeAIEmbeddings.
    from services.embeddings import GeminiHTTPEmbeddings
    return GeminiHTTPEmbeddings(model="gemini-embedding-001", output_dimensionality=768)

# --- UTILITIES ---
def extract_json_from_text(text: str) -> str:
    """Extracts JSON from an LLM response, stripping out markdown formatting."""
    if not text:
        return "{}"

    # Clean up standard markdown JSON code blocks
    text = text.replace('```json', '').replace('```', '').strip()

    start_idx = text.find('{')
    end_idx = text.rfind('}')

    if start_idx != -1 and end_idx != -1 and end_idx >= start_idx:
        return text[start_idx:end_idx + 1]

    return "{}" # Return empty JSON object if nothing found


def _repair_json_string(s: str) -> str:
    """Best-effort repair of common LLM JSON mistakes:
      - trailing commas before `}` / `]`
      - // line comments and /* block comments
      - smart quotes
      - unquoted object keys (e.g. `{ field: "x" }` → `{ "field": "x" }`)

    Conservative: only intended for fallback when strict json.loads fails."""
    import re as _re

    # Strip JS-style comments (LLMs sometimes annotate fields).
    s = _re.sub(r"//[^\n]*", "", s)
    s = _re.sub(r"/\*.*?\*/", "", s, flags=_re.DOTALL)

    # Normalise smart quotes.
    s = (s.replace("“", '"').replace("”", '"')
           .replace("‘", "'").replace("’", "'"))

    # Quote unquoted keys: matches `{` or `,` followed by an identifier then `:`.
    s = _re.sub(r'([{,]\s*)([A-Za-z_][A-Za-z0-9_]*)\s*:', r'\1"\2":', s)

    # Drop trailing commas before } or ].
    s = _re.sub(r",(\s*[}\]])", r"\1", s)

    return s


def parse_llm_json(text: str) -> dict:
    """Robust LLM-output → dict. Tries strict parse first, then a repair pass."""
    raw = extract_json_from_text(text)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as first_err:
        repaired = _repair_json_string(raw)
        try:
            return json.loads(repaired)
        except json.JSONDecodeError:
            # Re-raise the original (more informative) error.
            raise first_err

async def extract_text_from_file(base64_data: str, media_type: str) -> str:
    """Extracts text from PDF, Images, or Word documents with high-fidelity formatting."""
    
    # 1. Handle Word Documents (.docx)
    if "word" in media_type or "officedocument" in media_type:
        try:
            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            return "Error: Could not read Word document."

    # 2. Handle PDFs and Images via Gemini Vision OCR
    vision_llm = get_vision_llm()
    
    # ENHANCED EXTRACTION PROMPT
    vision_prompt = """
    You are a world-class medical document transcriptionist. 
    Your goal is to extract clinical data from the provided image/document and format it for EXTREME READABILITY.

    ### FORMATTING RULES (STRICT):
    1. MAIN TITLES: Put category names in bold followed by a colon (e.g., **Patient Identification:**).
    2. DATA FIELDS: List sub-titles (labels) followed by their info on separate lines under the main title.
    3. LINE BREAKS: Use a double line break between categories and a single line break between every data field.
       Example:
       **Patient Identification:**
       Medicare Coverage: Yes
       Patient Name: XX YY.

       **Insured Information:**
       Insured's I.D. Number: 987 65 4321A
    
    4. CLEAN UP: IGNORE all boilerplate instructions, footer codes, and "Leave Blank" fields.
    5. NO BOILERPLATE: Do not include text like "DO NOT WRITE IN THIS SPACE" or generic form titles.

    If the document is a blank form with no handwritten or typed data, output exactly: "[NO CLINICAL DATA FOUND - BLANK FORM]"
    """
    
    messages = [
        HumanMessage(content=[
            {"type": "text", "text": vision_prompt},
            {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{base64_data}"}}
        ])
    ]
    response = await vision_llm.ainvoke(messages)
    return response.content

METADATA_EXTRACTION_PROMPT = """You are a medical claim metadata extractor. Read the clinical documents below and extract the following fields. If a field is not clearly present in the documents, return null for that field — do NOT guess.

Fields to extract:
- patient_name: The patient's full name as it appears in the documents.
- patient_id: The patient's national ID, insurance ID, or medical record number.
- patient_age: Integer age in years.
- patient_gender: "Male", "Female", or null.
- patient_state: The patient's state, governorate, or region of residence (e.g. "Amman", "Dubai", "Riyadh"). Null if not stated.
- provider_name: The hospital, clinic, or provider organisation name.
- provider_specialty: The treating physician's specialty or the department issuing the request (e.g. "Cardiology", "Orthopedics", "General Surgery"). Null if not stated.
- diagnosis_code: ICD-10 code if present (e.g. "M54.5"), otherwise null.
- procedure_code: CPT or local procedure code if present, otherwise null.
- visit_type: "Inpatient", "Outpatient", "Emergency", "Day Surgery", or similar — whichever the document indicates. Null if unclear.
- length_of_stay: Numeric expected or actual length of stay in days. For purely outpatient visits, return 1. Null if not stated and not inferable.
- insurance_type: The plan, scheme, or coverage type referenced in the documents (e.g. "Comprehensive", "Basic", "Government", "Corporate"). Null if not stated.
- claim_amount: Numeric requested/billed amount (no currency), otherwise null.

If multiple documents disagree (e.g. two different patient names), pick the value that appears in the most documents and add a note in `extraction_notes`.

Respond ONLY with valid JSON in this exact shape:
{
    "patient_name": string | null,
    "patient_id": string | null,
    "patient_age": integer | null,
    "patient_gender": "Male" | "Female" | null,
    "patient_state": string | null,
    "provider_name": string | null,
    "provider_specialty": string | null,
    "diagnosis_code": string | null,
    "procedure_code": string | null,
    "visit_type": string | null,
    "length_of_stay": number | null,
    "insurance_type": string | null,
    "claim_amount": number | null,
    "extraction_notes": string | null
}
"""


async def extract_metadata_from_docs(combined_text: str) -> dict:
    """Extracts structured claim metadata from OCR'd document text. Returns a dict
    with whichever fields could be confidently identified; missing fields are absent."""
    if not combined_text or not combined_text.strip():
        return {}
    llm = get_llm()
    messages = [
        SystemMessage(content=METADATA_EXTRACTION_PROMPT),
        HumanMessage(content=combined_text[:12000]),
    ]
    try:
        response = await llm.ainvoke(messages)
        parsed = _safe_parse_json(response.content)
    except Exception as e:
        logger.error(f"Metadata extraction failed: {e}")
        return {}

    out = {}
    for key in (
        "patient_name", "patient_id",
        "patient_age", "patient_gender", "patient_state",
        "provider_name", "provider_specialty",
        "diagnosis_code", "procedure_code",
        "visit_type", "length_of_stay", "insurance_type",
        "claim_amount",
    ):
        val = parsed.get(key)
        if val is None:
            continue
        if isinstance(val, str) and not val.strip():
            continue
        out[key] = val
    if parsed.get("extraction_notes"):
        out["extraction_notes"] = parsed["extraction_notes"]
    return out


async def process_pre_auth_case(pre_auth_id: str, insurer_id: str, attachments: list):
    """Background task: OCR the dropped-off documents and extract structured
    metadata into any blank pre-auth columns. No AI decision or recommendation
    is produced for pre-auth — every routed request goes straight to the
    insurer's manual review queue."""
    logger.info(f"Starting background processing for Pre-Auth: {pre_auth_id}")

    # 1. Perform OCR on all documents
    for att in attachments:
        try:
            # att is a dict with file_name, content_type, content (base64)
            extracted_text = await extract_text_from_file(att["content"], att["content_type"])

            # Update the existing document record with extracted text
            supabase.table("pre_auth_documents").update({
                "extracted_text": extracted_text
            }).eq("pre_auth_id", pre_auth_id).eq("file_name", att["file_name"]).execute()

        except Exception as e:
            logger.error(f"Background OCR failed for {att['file_name']}: {e}")

    # 2. Pull all OCR'd text back and extract structured metadata from the docs.
    #    This replaces the patient/provider fields the provider used to type into
    #    the dropoff form — we now derive them from the documents themselves.
    docs_res = supabase.table("pre_auth_documents").select(
        "file_name, extracted_text"
    ).eq("pre_auth_id", pre_auth_id).execute()

    if docs_res.data:
        combined = "\n\n".join(
            f"--- {d['file_name']} ---\n{d.get('extracted_text') or ''}"
            for d in docs_res.data
        )
        metadata = await extract_metadata_from_docs(combined)
        if metadata:
            extracted = {k: v for k, v in metadata.items() if k != "extraction_notes"}
            # The submitter may have typed a structured packet into the form.
            # Document extraction only fills columns they left blank — it must
            # never clobber data the doctor/provider entered by hand.
            current = {}
            try:
                cur_res = supabase.table("pre_auth_requests").select("*").eq(
                    "id", pre_auth_id
                ).execute()
                if cur_res.data:
                    current = cur_res.data[0]
            except Exception as e:
                logger.warning(f"Pre-Auth {pre_auth_id}: could not read current row: {e}")

            def _is_blank(field: str, value) -> bool:
                if value is None:
                    return True
                if isinstance(value, str):
                    return not value.strip() or value.strip().lower() == "pending extraction"
                if field == "claim_amount":
                    return float(value) == 0.0
                return False

            updates = {
                k: v for k, v in extracted.items()
                if _is_blank(k, current.get(k))
            }
            skipped = [k for k in extracted if k not in updates]
            if skipped:
                logger.info(
                    f"Pre-Auth {pre_auth_id}: kept submitter-entered fields {skipped}"
                )
            if updates:
                logger.info(
                    f"Pre-Auth {pre_auth_id}: extracted metadata fields "
                    f"{list(updates.keys())}"
                )
                try:
                    supabase.table("pre_auth_requests").update(updates).eq(
                        "id", pre_auth_id
                    ).execute()
                except Exception as e:
                    logger.error(f"Failed to persist extracted metadata: {e}")

    logger.info(
        f"Pre-Auth {pre_auth_id}: documents OCR'd and metadata extracted. "
        "No AI review runs on pre-auth — it awaits the insurer's manual decision."
    )

# --- CORE SERVICES ---
def _safe_parse_json(raw: str) -> dict:
    try:
        parsed = json.loads(extract_json_from_text(raw))
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


async def process_and_embed_policy_file(insurer_id: str, base64_data: str, file_name: str = ""):
    """Called once when an insurer uploads their policy. Chunks the PDF/Word and saves to Supabase."""
    logger.info(f"Starting policy file processing for insurer {insurer_id}")
    
    full_text = ""
    is_word = file_name.lower().endswith(".docx") or "officedocument" in base64_data[:50] 

    try:
        if is_word:
            # Handle Word
            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            # Handle PDF (Default)
            pdf_bytes = base64.b64decode(base64_data)
            pdf = pdfium.PdfDocument(pdf_bytes)
            for page in pdf:
                text_page = page.get_textpage()
                full_text += text_page.get_text_range() + "\n\n"
    except Exception as e:
        logger.error(f"Failed to parse policy file: {e}")
        raise ValueError(f"Could not read the uploaded file. Please ensure it is a valid PDF or .docx file. Error: {str(e)}")
        
    if not full_text.strip():
        logger.warning(f"Insurer {insurer_id} uploaded a file with no readable text.")
        raise ValueError("The uploaded file contains no readable text.")

    # 2. Chunk text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_text(full_text)
    logger.info(f"Split PDF into {len(chunks)} chunks. Beginning embedding...")

    embeddings_model = get_embeddings()
    
    # 3. Clear old policy chunks for this insurer
    supabase.table("policy_chunks").delete().eq("insurer_id", insurer_id).execute()
    
    # 4. Embed and insert in safe batches to respect rate limits
    batch_size = 50 
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i+batch_size]
        
        try:
            vectors = embeddings_model.embed_documents(batch_chunks)
        except Exception as e:
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                logger.warning("Rate limit hit during embedding. Retrying in 10 seconds...")
                await asyncio.sleep(10) 
                vectors = embeddings_model.embed_documents(batch_chunks)
            else:
                raise e
        
        payload = []
        for chunk, vector in zip(batch_chunks, vectors):
            payload.append({
                "insurer_id": insurer_id,
                "content": chunk,
                "embedding": vector
            })
            
        supabase.table("policy_chunks").insert(payload).execute()
        logger.info(f"Inserted batch {i//batch_size + 1} into database.")
        
        await asyncio.sleep(3)
        
    logger.info(f"Finished processing policy for insurer {insurer_id}. Total {len(chunks)} chunks saved.")


# ============================================================================
# PROVIDER-SIDE AI SERVICES
# Used by /api/claims and /api/intake routers for claim extraction & scrubbing.
# ============================================================================

from typing import List
from pydantic import BaseModel as _BaseModel, Field as _Field
from langchain_core.output_parsers import PydanticOutputParser


class FieldString(_BaseModel):
    value: str = _Field(default="", description="Extracted text. Empty string if not found.")
    confidence: int = _Field(default=0, description="Confidence 0-100. 0 if not found.")


class FieldFloat(_BaseModel):
    value: float = _Field(default=0.0)
    confidence: int = _Field(default=0)


class FieldStringList(_BaseModel):
    value: List[str] = _Field(default_factory=list)
    confidence: int = _Field(default=0)


class ClaimExtractionResult(_BaseModel):
    patient_name: FieldString
    patient_id: FieldString
    date_of_service: FieldString = _Field(description="Format YYYY-MM-DD")
    provider_name: FieldString
    provider_id: FieldString
    payer_name: FieldString
    member_id: FieldString
    primary_diagnosis: FieldString = _Field(description="Primary ICD-10 code")
    additional_diagnoses: FieldStringList = _Field(description="Additional ICD-10 codes")
    primary_procedure: FieldString = _Field(description="Primary CPT/HCPCS code")
    additional_procedures: FieldStringList = _Field(description="Additional CPT/HCPCS codes")
    billed_amount: FieldFloat
    additional_notes: FieldString
    # Fraud-detector signals — extracted only if visibly present in the document.
    patient_age: FieldString = _Field(description="Patient age as a whole number — '' if not stated.")
    patient_gender: FieldString = _Field(description="'Male' or 'Female' — '' if not stated.")
    patient_state: FieldString = _Field(description="Patient state/governorate/region (e.g. 'Amman'). '' if not stated.")
    visit_type: FieldString = _Field(description="One of: Inpatient, Outpatient, Emergency, Day Surgery — '' if not stated.")
    length_of_stay: FieldString = _Field(description="Length of stay in days (whole number). '' if not stated or outpatient.")
    insurance_type: FieldString = _Field(description="Plan type referenced on the document (e.g. 'Comprehensive', 'Basic'). '' if not stated.")
    provider_specialty: FieldString = _Field(description="Treating physician specialty (e.g. 'Cardiology'). '' if not stated.")


_claim_parser = PydanticOutputParser(pydantic_object=ClaimExtractionResult)


CLAIM_EXTRACTION_PROMPT = """You are an expert medical claims parser.
Extract the data directly from the provided image.

CRITICAL RULES:
1. NO HALLUCINATIONS: If a field is NOT clearly present in the image, set its value to "" (or 0) and its confidence to 0. Do not invent notes, IDs, demographics, or clinical context.
2. CONFIDENCE SCORES: Rate your confidence from 0 to 100 based on how clearly you can read the text.
3. Normalize dates to YYYY-MM-DD.
4. For the fraud-signal fields (patient_age, patient_gender, patient_state, visit_type, length_of_stay, insurance_type, provider_specialty): extract ONLY when explicitly stated in the document. Do not guess from context. Empty + confidence 0 is preferable to a guess.
5. visit_type must be exactly one of: Inpatient, Outpatient, Emergency, Day Surgery. Map common synonyms (e.g. "ambulatory" → Outpatient, "ER" → Emergency) but leave blank if unclear.

{format_instructions}
"""


SCRUB_SYSTEM_PROMPT = """You are ClaimRidge AI, an expert medical claims scrubbing engine. The provider has ALREADY delivered care and is now billing. Your job is RETROSPECTIVE billing review: catch coding errors, verify the procedure-diagnosis link is supported, check for unbundling/upcoding patterns, and confirm authorization was in place when required.

You are NOT a medical-necessity gatekeeper — necessity was decided at pre-authorisation time. Your focus is whether the bill in front of you is coded correctly and consistent with the auth that authorised it.

{unregistered_warning}

## Authorization status for this claim
<auth_check>
{auth_check_summary}
</auth_check>

If the auth status is `missing`, `not_approved`, `expired`, `wrong_patient`, or `code_mismatch`, you MUST raise an `error`-severity issue on the `pre_auth_number` field with severity `error` and a clear message explaining the problem. This is non-negotiable — most payers deny claims with broken authorization linkage.

If the auth block says **NOT VERIFIABLE** (out-of-network payer), do NOT raise any issue on `pre_auth_number`. We have no way to verify it against an external payer's system, so accept the supplied number as informational metadata and move on.

If **NOT PROVIDED** (no number supplied at all), do not raise an error either — instead note in `recommendations` that the provider should confirm whether an auth was required for these procedures.

## Coding review checklist (apply in order)
1. **CPT ↔ ICD-10 alignment** — Does at least one billed procedure code have a clinically supported diagnosis code on the claim? Flag procedures with no supporting diagnosis.
2. **Bundling / NCCI edits** — Are billed procedures normally bundled into another billed code? (e.g., minor procedures inside a major one). If yes, suggest the bundled code.
3. **Upcoding patterns** — Is there a higher-tier code billed where the documented procedure description matches a lower-tier code? Flag as `warning`.
4. **Modifier requirements** — Do bilateral / multiple / repeat procedures have appropriate modifiers? (If the codebase has modifier data, comment on missing ones.)
5. **Units sanity** — Are units reasonable for the procedure type? (e.g., a single surgery shouldn't be billed with units > 1 unless explicitly bilateral.)
6. **Fee schedule realism** — Is the `billed_amount` in a reasonable range for the billed codes? Extreme outliers → `warning`.
7. **Place of service vs setting** — Does the procedure type match the inpatient/outpatient setting suggested by other claim fields?
8. **Required documentation hint** — If the payer policy explicitly requires a specific attached document (operative note, anaesthesia record), call it out in `recommendations`.

## Payer-Specific Policy Rules (CRITICAL)
Use these to find PAYER-SPECIFIC requirements (specific modifier rules, bundling exceptions, documentation requirements). DO NOT use them to second-guess medical necessity — that was already decided.
<payer_rules>
{policy_context}
</payer_rules>

## Response Format
Return ONLY valid JSON with this exact structure:
{{
  "status": "clean" | "warnings" | "errors",
  "overall_score": <number 0-100, where 100 = clean clean claim>,
  "issues": [
    {{
      "field": "<exact field name — e.g. procedure_codes[0], pre_auth_number, billed_amount>",
      "severity": "error" | "warning" | "info",
      "message": "<specific problem in coding/billing terms>",
      "suggestion": "<exact fix the biller should apply>"
    }}
  ],
  "corrected_claim": {{ <corrected version of the input> }},
  "recommendations": ["<actionable recommendations focused on billing accuracy and auth compliance>"],
  "applied_policy_rules": ["<Quote the exact rules from the payer_rules section that you actually used>"]
}}
"""


# The final claim adjudicator. Runs AFTER fraud screening and the coding
# scrubber, only for claims that already cleared the fraud hard-gate (i.e. NOT
# high/extreme risk) and whose payer has a policy on file. It issues the final
# automatic verdict: accept | deny | escalate.
ADJUDICATION_SYSTEM_PROMPT = """You are ClaimRidge AI, the final claims adjudicator for a health insurer in the MENA region. The provider has ALREADY delivered care and submitted the bill. This claim has already cleared statistical fraud screening (it is NOT high-risk) and has been through a coding scrubber. Your job is to issue the FINAL automatic decision: accept it for payment, deny it, or escalate it to a human medical reviewer.

You are adjudicating BILLING and POLICY COMPLIANCE — not medical necessity, which was settled at pre-authorisation. Decide whether this bill should be paid as submitted.

## Inputs you are given
- The structured claim (patient, codes, amount, pre-auth linkage).
- The coding scrubber's findings — issues it already detected.
- The authorization-linkage verdict for any referenced pre-authorisation.
- The payer's own policy rules, retrieved below.

## Authorization linkage verdict
<auth_check>
{auth_check_summary}
</auth_check>

## Coding scrubber findings
<scrub_findings>
{scrub_findings}
</scrub_findings>

## Payer policy rules (RAG context)
<payer_rules>
{policy_context}
</payer_rules>

## DECISION — choose ONE

- **accept**: The claim is coded correctly, the billed amount is reasonable for the procedures, authorization (if referenced) is verified or was not required, and no payer policy rule is violated. Accepting means the claim is automatically approved for payment with no human review.
- **deny**: There is a CONCRETE, denial-grade defect that payers routinely deny for. Valid grounds: a broken authorization linkage (the auth verdict is MISSING, NOT APPROVED, EXPIRED, WRONG PATIENT, or CODE MISMATCH); a billed procedure with no supporting diagnosis on the claim; an explicit payer policy exclusion that applies; or a clear billing-integrity violation (e.g. a procedure unbundled against an NCCI edit). You MUST cite the specific defect.
- **escalate**: You see a real concern that is not clearly denial-grade and needs a human medical reviewer — conflicting or ambiguous documentation, a coding warning whose severity you cannot resolve, or a policy rule whose application is unclear. You MUST cite what the reviewer needs to resolve.

## CONSTRAINTS
- `deny` and `escalate` require CONCRETE evidence: a specific code, field, policy rule, or auth-verdict value. Vague phrases like "needs further review" are NOT evidence and will be rejected.
- A broken authorization verdict (MISSING / NOT APPROVED / EXPIRED / WRONG PATIENT / CODE MISMATCH) is denial-grade — choose `deny`, not `escalate`.
- If the auth block says NOT VERIFIABLE (out-of-network) or NOT PROVIDED, do NOT treat authorization as a defect.
- Do NOT deny on suspicion of fraud — fraud screening already cleared this claim. Deny only on coding, billing, policy, or authorization grounds.
- Do NOT invent policy rules that are not in the payer rules section.
- When genuinely unsure between accept and escalate, escalate. When the defect is concrete and denial-grade, deny.

## TONE
The rationale is shown to the insurer's staff and the provider. Write in clear, professional plain English. Do not mention "model", "score", or internal mechanics.

## OUTPUT
Respond ONLY with valid JSON, no commentary:
{{
  "decision": "accept" | "deny" | "escalate",
  "rationale": "A clear, professional paragraph (3-5 sentences) explaining the decision.",
  "evidence": [
    {{
      "step": "coding_review" | "authorization" | "policy_compliance" | "billing_integrity",
      "finding": "A specific, professionally worded observation naming the exact code, field, policy rule, or authorization verdict."
    }}
  ],
  "policy_basis": ["Quote the exact payer rules you relied on, if any."]
}}
"""


# --- AI MEDICAL NECESSITY REVIEW (claim-side, advisory) ---------------------
# Surfaced on the insurer claim-detail page as "AI Medical Necessity Review"
# (the "Generate Clinical Review" action). It is ADVISORY clinical input only —
# it never decides accept/deny. The binding verdict is owned by claim
# adjudication (services/adjudication.py) and the human medical officer.
MEDICAL_NECESSITY_PROMPT = """You are ClaimRidge AI, a clinical reviewer producing an ADVISORY medical-necessity opinion on a submitted insurance claim for the MENA market (Jordan / GCC / Levant).

This claim has already been billed — the service was delivered. Your job is NOT to accept or deny it: a separate adjudication step and a human medical officer own that decision. Your job is to give the reviewer a clear, evidence-based read on whether the documented clinical picture justifies the billed care under THIS payer's policy.

## WHAT YOU ARE GIVEN
- The billed diagnosis (ICD-10) and procedure (CPT) codes, each with its catalogue description.
- The provider's clinical notes and any OCR'd text from clinical documents attached to a linked pre-authorisation.
- The payer's own policy excerpts retrieved for this claim (may be empty if the payer has no policy on file).
- Signals already computed by the system: coding-scrubber issues, the pre-auth authorisation check, and the fraud-screening result. Treat these as context, not as your verdict.

## HOW TO REVIEW — assess each criterion
1. **Clinical indication** — do the documented findings, symptoms, or history justify performing this procedure for this diagnosis?
2. **Severity & setting of care** — is the intensity/setting of the billed care (outpatient / day-surgery / inpatient) proportional to the documented severity?
3. **Conservative management** — where evidence-based practice or the payer policy expects less-invasive options first, is that trial documented?
4. **Diagnosis-procedure coherence** — does the procedure follow logically from the diagnosis, with no unexplained or unrelated billing?
5. **Payer policy alignment** — does the claim meet the specific coverage criteria in the payer policy excerpts? If no policy excerpts were provided, judge against standard MCG/InterQual-style and WHO/NICE guidance and say so.

For each criterion return a status:
- `met` — the documentation clearly supports it.
- `partial` — partially supported; a specific gap remains.
- `not_met` — the documentation contradicts it or a required element is clearly absent.
- `unknown` — the supplied documentation is insufficient to judge this criterion.

## OVERALL ASSESSMENT — choose ONE
- `supported` — documentation supports medical necessity for the billed care.
- `partially_supported` — necessity is supported for part of the claim; specific gaps or weaker elements remain.
- `insufficient_documentation` — too little clinical information was supplied to reach a view; name what is missing.
- `not_supported` — the documented clinical picture does not justify the billed care.

## EVIDENCE STANDARD
- Every `finding` must be CONCRETE: name a specific documented finding, a missing element, a code, a policy criterion, or quote the notes. Vague phrases ("further review needed", "appears appropriate") are not findings.
- Do not invent policy requirements that are not in the supplied excerpts.
- `policy_basis` must quote or closely paraphrase the supplied payer policy excerpts only. Leave it empty if none were supplied.
- `recommended_actions` are advisory next steps for the human reviewer (e.g. "request the operative report", "confirm the conservative-care trial dates"). Keep them specific and actionable.

## TONE
The output is read by the insurer's medical reviewer. Write in clear, professional, plain clinical English. Do not mention internal mechanics (no "model", "score", "fraud layer", "prompt", percentages, or debug language). Frame concerns as documentation observations, not accusations.

## OUTPUT
Respond ONLY with valid JSON, no commentary, in this exact shape:
{{
    "assessment": "supported" | "partially_supported" | "insufficient_documentation" | "not_supported",
    "headline": "A single advisory sentence summarising the medical-necessity read.",
    "summary": "2-3 sentence clinical reasoning paragraph in plain English.",
    "criteria": [
        {{
            "name": "Clinical indication" | "Severity & setting of care" | "Conservative management" | "Diagnosis-procedure coherence" | "Payer policy alignment",
            "status": "met" | "partial" | "not_met" | "unknown",
            "finding": "A specific, professionally worded observation citing the documentation."
        }}
    ],
    "policy_basis": [
        {{ "snippet": "Quoted or closely paraphrased payer policy text.", "note": "How it applies to this claim." }}
    ],
    "recommended_actions": [ "A specific advisory next step for the reviewer." ]
}}
"""

ALLOWED_MN_ASSESSMENTS = {
    "supported", "partially_supported", "insufficient_documentation", "not_supported",
}
MN_CRITERION_STATUSES = {"met", "partial", "not_met", "unknown"}


def _validate_medical_necessity(result: dict) -> tuple[bool, str]:
    """Returns (is_valid, reason) for a structured medical-necessity result."""
    if not isinstance(result, dict):
        return False, "response is not a JSON object"
    assessment = str(result.get("assessment", "")).strip().lower()
    if assessment not in ALLOWED_MN_ASSESSMENTS:
        return False, f"assessment must be one of {sorted(ALLOWED_MN_ASSESSMENTS)}, got '{assessment}'"
    if not str(result.get("summary", "")).strip():
        return False, "summary is empty"
    criteria = result.get("criteria")
    if not isinstance(criteria, list) or not criteria:
        return False, "criteria must be a non-empty array"
    for item in criteria:
        if not isinstance(item, dict):
            continue
        status = str(item.get("status", "")).strip().lower()
        finding = str(item.get("finding", "")).strip()
        if status in MN_CRITERION_STATUSES and len(finding) >= 20:
            return True, ""
    return False, "no criterion has a valid status and a concrete finding"


def _medical_necessity_fallback() -> dict:
    """Advisory result used when the LLM cannot produce a valid structured review.
    Frames it as 'review manually' rather than guessing a clinical verdict."""
    return {
        "assessment": "insufficient_documentation",
        "headline": "An advisory clinical review could not be completed automatically.",
        "summary": (
            "The automated medical-necessity review could not reach a confident "
            "conclusion for this claim. A medical reviewer should assess the billed "
            "diagnosis and procedure against the clinical documentation directly."
        ),
        "criteria": [],
        "policy_basis": [],
        "recommended_actions": [
            "Open the supporting documents and review medical necessity manually.",
        ],
        "policy_available": False,
    }


async def extract_claim_from_documents(documents: list[dict]) -> dict:
    """Multi-document claim extraction. Each entry in `documents` must be
    {"fileBase64": str, "mediaType": str, "fileName": str}. All images are sent
    in one Gemini call so the model can cross-reference fields across them
    (e.g. patient name from an insurance card + diagnosis from a clinical
    note) and return a single consolidated record."""
    if not documents:
        raise ValueError("No documents provided.")

    vision_llm = get_vision_llm()
    prompt_text = CLAIM_EXTRACTION_PROMPT.format(
        format_instructions=_claim_parser.get_format_instructions()
    )

    content: list[dict] = [{"type": "text", "text": prompt_text}]
    if len(documents) > 1:
        content.append({
            "type": "text",
            "text": (
                f"You are looking at {len(documents)} related documents for the SAME claim "
                "(e.g. claim form, insurance card, clinical note, lab report). "
                "Cross-reference fields across them and pick the single most reliable value for each field. "
                "Confidence should reflect agreement across documents — conflicting values lower confidence."
            ),
        })
    for d in documents:
        file_name = d.get("fileName") or "document"
        media_type = d["mediaType"]
        file_b64 = d["fileBase64"]
        content.append({"type": "text", "text": f"--- Document: {file_name} ---"})
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{media_type};base64,{file_b64}"},
        })

    messages = [HumanMessage(content=content)]
    logger.info(
        f"Calling Vision Model ({Config.OCR_MODEL}) for claim extraction "
        f"with {len(documents)} document(s)..."
    )
    try:
        response = await vision_llm.ainvoke(messages)
        result = _claim_parser.parse(response.content)
        return result.model_dump()
    except Exception as e:
        logger.error(f"Multi-doc claim extraction failed: {str(e)}")
        raise ValueError("The AI failed to format the document properly. Please re-upload.") from e


async def extract_claim_from_document(file_base64: str, media_type: str) -> dict:
    """Backward-compat shim around extract_claim_from_documents for single-file callers."""
    return await extract_claim_from_documents([
        {"fileBase64": file_base64, "mediaType": media_type, "fileName": "document"}
    ])


# ----------------------------------------------------------------------------
# Pre-authorisation document extraction
# Mirrors the claim extractor, but pulls the prospective pre-auth packet so the
# pre-auth form can be auto-filled before the provider/doctor reviews it.
# ----------------------------------------------------------------------------
class ExtractedCode(_BaseModel):
    code: str = _Field(default="", description="The code exactly as written in the document.")
    description: str = _Field(
        default="",
        description="The description/label printed next to the code in the document. '' if none is shown.",
    )


class FieldCodeList(_BaseModel):
    value: List[ExtractedCode] = _Field(default_factory=list)
    confidence: int = _Field(default=0, description="Confidence 0-100.")


class PreAuthExtractionResult(_BaseModel):
    patient_name: FieldString
    patient_dob: FieldString = _Field(description="Date of birth, format YYYY-MM-DD.")
    patient_gender: FieldString = _Field(description="'Male' or 'Female' — '' if not stated.")
    patient_id: FieldString = _Field(description="National / patient ID number.")
    insurance_member_id: FieldString = _Field(description="Health insurance member/policy ID.")
    insurance_group_number: FieldString
    patient_phone: FieldString
    patient_address: FieldString
    ordering_provider_name: FieldString = _Field(description="Doctor who ordered/requested the treatment.")
    ordering_provider_npi: FieldString
    ordering_provider_tax_id: FieldString
    servicing_provider_name: FieldString = _Field(description="Facility/doctor performing the procedure — '' if same as ordering.")
    servicing_provider_npi: FieldString
    servicing_provider_tax_id: FieldString
    diagnosis_codes: FieldCodeList = _Field(description="All ICD-10 diagnosis codes visible, each with its description.")
    procedure_codes: FieldCodeList = _Field(description="All CPT/HCPCS procedure codes visible, each with its description.")
    modifiers: FieldString = _Field(description="CPT modifiers (e.g. LT, RT, 59).")
    ndc_code: FieldString = _Field(description="NDC drug code for specialty pharmacy requests.")
    place_of_service: FieldString = _Field(description="Where care happens — see allowed list in the prompt.")
    anticipated_date_of_service: FieldString = _Field(description="Planned service date, YYYY-MM-DD.")
    payer_name: FieldString = _Field(description="Insurance company / payer the request is addressed to.")


_pre_auth_parser = PydanticOutputParser(pydantic_object=PreAuthExtractionResult)


PRE_AUTH_EXTRACTION_PROMPT = """You are an expert medical pre-authorisation parser.
A pre-authorisation is a PROSPECTIVE request — the provider asks the insurer to greenlight a planned procedure before it happens. Extract the request packet directly from the provided document image(s).

CRITICAL RULES:
1. NO HALLUCINATIONS: If a field is NOT clearly present, set its value to "" (or an empty list) and its confidence to 0. Never invent IDs, demographics, codes, or provider details.
2. CONFIDENCE SCORES: Rate confidence 0-100 based on how clearly you can read each field.
3. Normalize every date to YYYY-MM-DD.
4. patient_gender must be exactly "Male" or "Female" — '' if unclear.
5. place_of_service must be one of: Inpatient Hospital, Outpatient Hospital, Outpatient Surgery Center, Doctor's Office, Patient's Home, Emergency Room, Telehealth. Map synonyms (e.g. "ambulatory surgery centre" -> Outpatient Surgery Center, "ER" -> Emergency Room). Leave '' if unclear.
6. diagnosis_codes = ICD-10 codes; procedure_codes = CPT/HCPCS codes. Capture EVERY code shown. For each code, also copy the description/label printed beside it in the document into its `description` field; if the document prints no description, set it to "". If a condition or procedure is written out in words with no code number at all, leave it out.
7. The ordering provider decided the treatment is necessary; the servicing provider performs it. Only fill the servicing_* fields if the document clearly distinguishes a different performing provider/facility.
8. payer_name = the insurance company the request is addressed to.

{format_instructions}
"""


async def extract_pre_auth_from_documents(documents: list[dict]) -> dict:
    """Multi-document pre-auth extraction. Each entry must be
    {"fileBase64": str, "mediaType": str, "fileName": str}. All images are sent
    in one Gemini call so the model can cross-reference fields and return a
    single consolidated pre-auth packet with per-field confidence scores."""
    if not documents:
        raise ValueError("No documents provided.")

    vision_llm = get_vision_llm()
    prompt_text = PRE_AUTH_EXTRACTION_PROMPT.format(
        format_instructions=_pre_auth_parser.get_format_instructions()
    )

    content: list[dict] = [{"type": "text", "text": prompt_text}]
    if len(documents) > 1:
        content.append({
            "type": "text",
            "text": (
                f"You are looking at {len(documents)} related documents for the SAME "
                "pre-authorisation request. Cross-reference fields across them and pick "
                "the single most reliable value for each. Conflicting values lower confidence."
            ),
        })
    for d in documents:
        file_name = d.get("fileName") or "document"
        media_type = d["mediaType"]
        file_b64 = d["fileBase64"]
        content.append({"type": "text", "text": f"--- Document: {file_name} ---"})
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{media_type};base64,{file_b64}"},
        })

    messages = [HumanMessage(content=content)]
    logger.info(
        f"Calling Vision Model ({Config.OCR_MODEL}) for pre-auth extraction "
        f"with {len(documents)} document(s)..."
    )
    try:
        response = await vision_llm.ainvoke(messages)
        result = _pre_auth_parser.parse(response.content)
        return result.model_dump()
    except Exception as e:
        logger.error(f"Multi-doc pre-auth extraction failed: {str(e)}")
        raise ValueError("The AI failed to read the document(s) properly. Please re-upload.") from e


def _format_auth_check(auth_check: dict | None) -> str:
    """Turns the auth_check verdict into a paragraph the LLM can read."""
    if not auth_check:
        return "No authorization check was performed."
    status = (auth_check.get("status") or "").lower()
    detail = auth_check.get("detail") or ""
    # `not_applicable` covers two distinct cases — disambiguate via the detail
    # so the prompt's downstream rules trigger correctly.
    if status == "not_applicable":
        if "out-of-network" in detail.lower():
            head = "NOT VERIFIABLE — The payer is out-of-network. Treat the supplied authorization number as opaque metadata; do NOT flag it."
        else:
            head = "NOT PROVIDED — No pre-authorisation number was supplied with this claim."
        return f"{head}\n{detail}".strip()
    labels = {
        "ok": "VERIFIED — A valid pre-authorisation covers this claim.",
        "missing": "MISSING — The provider referenced a pre-authorisation, but no matching request exists for this payer.",
        "not_approved": "NOT APPROVED — The referenced pre-authorisation exists but has not been approved by the insurer.",
        "expired": "EXPIRED — The referenced pre-authorisation is past its validity window.",
        "wrong_patient": "WRONG PATIENT — The pre-authorisation was issued for a different patient.",
        "code_mismatch": "CODE MISMATCH — The billed procedure codes are not within the authorised scope.",
        "contradiction": "CONTRADICTION — The claim's patient, coding, or provider details do not match the approved pre-authorisation.",
    }
    head = labels.get(status, f"STATUS: {status}")
    return f"{head}\n{detail}".strip()


# ---------------------------------------------------------------------------
# Deterministic auth severity. The scrubber prompt asks the LLM to flag broken
# authorization, but a 70B model can ignore that instruction in either
# direction — flagging an out-of-network number it was told to leave alone, or
# forgetting to flag a genuinely broken in-network auth. The auth verdict is
# already computed deterministically by services/authorization.py, so we
# *apply* it here in code rather than trusting the model.
# ---------------------------------------------------------------------------
_BROKEN_AUTH = {
    "missing", "not_approved", "expired", "wrong_patient", "code_mismatch",
    "contradiction",
}
# In-network claims with a broken auth verdict are denial-grade for most
# payers — cap the score so severity is guaranteed even if the LLM under-reacts.
_IN_NETWORK_BROKEN_AUTH_SCORE_CAP = 25


def _is_auth_issue(issue: dict) -> bool:
    """True if an LLM-produced issue is about pre-auth / authorization linkage."""
    if not isinstance(issue, dict):
        return False
    field = str(issue.get("field", "")).strip().lower()
    if field in {"pre_auth_number", "pre_auth", "authorization_number"}:
        return True
    text = f"{issue.get('message', '')} {issue.get('suggestion', '')}".lower()
    return any(kw in text for kw in ("pre-auth", "pre auth", "authorization", "authorisation"))


def _recompute_scrub_status(issues: list) -> str:
    """Derives the top-level status from the (possibly edited) issue list."""
    severities = {str(i.get("severity", "info")).lower() for i in issues if isinstance(i, dict)}
    if "error" in severities:
        return "errors"
    if "warning" in severities:
        return "warnings"
    return "clean"


def _enforce_auth_severity(result: dict, auth_check: dict | None, registered_payer_id: str | None) -> dict:
    """Applies the authorization verdict to the scrub result deterministically.

    - Out-of-network (payer not registered): the auth number cannot be verified
      against any payer we know, so it is NOT a defect. Strip any auth issue the
      LLM raised and add a single info-level "out-of-network — be careful"
      advisory. The score is left untouched — the number is withheld from the
      LLM upstream, so it never influenced scoring.
    - In-network with a broken auth verdict (missing / expired / wrong_patient /
      code_mismatch): guarantee an error-severity issue, force status to
      `errors`, and cap the score. Most payers deny these outright.
    """
    if not isinstance(result, dict):
        return result
    issues = result.get("issues")
    if not isinstance(issues, list):
        issues = []
    status = str((auth_check or {}).get("status", "")).lower()

    if not registered_payer_id:
        # Out-of-network — neglect the pre-auth, add a "be careful" advisory.
        issues = [i for i in issues if not _is_auth_issue(i)]
        issues.append({
            "field": "payer_name",
            "severity": "info",
            "message": (
                "This payer is out-of-network. ClaimRidge cannot verify the "
                "insurer's policy, coverage, or any pre-authorization for this claim."
            ),
            "suggestion": (
                "Confirm coverage and authorization requirements directly with the "
                "payer before submitting — this claim is recorded for your files "
                "only and is not validated against payer rules."
            ),
        })
        result["issues"] = issues
        result["status"] = _recompute_scrub_status(issues)
        return result

    if status in _BROKEN_AUTH:
        # In-network broken auth — be severe.
        detail = (auth_check or {}).get("detail") or (
            "The referenced pre-authorization is not valid for this claim."
        )
        if any(_is_auth_issue(i) for i in issues):
            for i in issues:
                if _is_auth_issue(i):
                    i["severity"] = "error"
        else:
            issues.append({
                "field": "pre_auth_number",
                "severity": "error",
                "message": detail,
                "suggestion": (
                    "Resolve the authorization before submitting — obtain a valid "
                    "pre-authorization covering this patient and the billed procedures."
                ),
            })
        result["issues"] = issues
        result["status"] = "errors"
        score = result.get("overall_score")
        if not isinstance(score, (int, float)) or score > _IN_NETWORK_BROKEN_AUTH_SCORE_CAP:
            result["overall_score"] = _IN_NETWORK_BROKEN_AUTH_SCORE_CAP
        return result

    return result


async def scrub_claim(claim_data: dict, registered_payer_id: str = None) -> dict:
    """Scrubs a structured claim against payer policy rules (RAG) and returns issues + corrections."""
    llm = get_llm(json_mode=True)

    # Auth check verdict is consumed by the prompt directly — strip it from the
    # JSON payload sent to the LLM so it doesn't confuse the coding review.
    auth_check = claim_data.pop("auth_check", None) if isinstance(claim_data, dict) else None
    auth_check_summary = _format_auth_check(auth_check)

    # Out-of-network: the pre-auth number can't be verified against any payer we
    # know, so it is not a coding defect. Withhold it from the LLM entirely so
    # the model can neither flag it nor let it drag down the coding score — the
    # out-of-network advisory is added back deterministically after scrubbing.
    if not registered_payer_id and isinstance(claim_data, dict):
        claim_data.pop("pre_auth_number", None)

    claim_json_str = json.dumps(claim_data, indent=2)

    unregistered_warning = ""
    policy_context = "No specific payer rules found. Use standard medical billing guidelines."
    retrieved_chunks = []

    if registered_payer_id:
        try:
            search_query = (
                f"Rules for Diagnoses: {', '.join(claim_data.get('diagnosis_codes', []))} "
                f"and Procedures: {', '.join(claim_data.get('procedure_codes', []))}"
            )
            embeddings_model = get_embeddings()
            query_vector = embeddings_model.embed_query(search_query)
            res = supabase.rpc("match_policy_rules", {
                "query_embedding": query_vector,
                "match_threshold": 0.4,
                "match_count": 4,
                "p_insurer_id": registered_payer_id,
            }).execute()
            if res.data:
                retrieved_chunks = [m["content"] for m in res.data]
                policy_context = "\n---\n".join(retrieved_chunks)
        except Exception as e:
            logger.error(f"RAG retrieval failed: {e}")
    else:
        unregistered_warning = (
            "NOTE: This claim's payer is not registered with ClaimRidge. "
            "Apply standard medical billing guidelines."
        )

    system_prompt = SCRUB_SYSTEM_PROMPT.format(
        unregistered_warning=unregistered_warning,
        policy_context=policy_context,
        auth_check_summary=auth_check_summary,
    )
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=f"Analyze and scrub the following medical claim:\n\n{claim_json_str}"),
    ]
    response = await llm.ainvoke(messages)
    try:
        result = parse_llm_json(response.content)
    except Exception as e:
        logger.error(f"Failed to parse scrub JSON: {str(e)} | raw: {response.content[:500]!r}")
        raise ValueError("AI returned invalid JSON during scrubbing.") from e

    result["retrieved_sources"] = retrieved_chunks
    # Auth severity is enforced in code, not left to the LLM (see above).
    result = _enforce_auth_severity(result, auth_check, registered_payer_id)
    return result


# ----------------------------------------------------------------------------
# Claim adjudication — the final accept/deny/escalate verdict
# Consumes the coding scrubber's findings + the fraud result + payer policy.
# Orchestrated by services/adjudication.py; this is only the LLM step.
# ----------------------------------------------------------------------------
ADJUDICATION_EVIDENCE_STEPS = {
    "coding_review", "authorization", "policy_compliance", "billing_integrity",
}

# Tokens that indicate a finding refers to something concrete in the claim or
# document (a field, a value, a contradiction, a quoted phrase) rather than
# meta-language about the review process itself.
CONCRETE_INDICATORS = {
    "name", "id", "age", "gender", "diagnosis", "procedure", "physician", "doctor",
    "patient", "provider", "policy rule", "amount", "date", "code", "field",
    "document", "form", "mismatch", "contradict", "missing", "absent", "blank",
    "quote", "states", "reads", "shows", "reports",
}

GENERIC_PHRASES = (
    "fraud system", "anti-fraud", "high-risk flag", "high risk flag",
    "further review", "further investigation", "necessitates review",
    "thorough examination", "due diligence",
)


def _is_concrete_finding(finding: str) -> bool:
    """True if a finding cites something specific (a field, value, code, quote,
    or contradiction) rather than vague review boilerplate."""
    if not finding or len(finding.strip()) < 30:
        return False
    f_lower = finding.lower()
    if any(ind in f_lower for ind in CONCRETE_INDICATORS):
        # Reject if it's *only* generic boilerplate even though it contains a keyword.
        non_generic = f_lower
        for phrase in GENERIC_PHRASES:
            non_generic = non_generic.replace(phrase, "")
        return len(non_generic.strip()) >= 30
    return False


def _validate_adjudication(result: dict) -> tuple[bool, str]:
    """Returns (is_valid, reason). `deny`/`escalate` must carry concrete evidence."""
    if not isinstance(result, dict):
        return False, "response is not a JSON object"
    decision = str(result.get("decision", "")).strip().lower()
    if decision not in {"accept", "deny", "escalate"}:
        return False, f"decision must be accept/deny/escalate, got '{decision}'"
    if decision == "accept":
        return True, ""
    evidence = result.get("evidence")
    if not isinstance(evidence, list) or not evidence:
        return False, f"{decision} decision requires a non-empty evidence array"
    for item in evidence:
        if not isinstance(item, dict):
            continue
        step = str(item.get("step", "")).strip().lower()
        finding = str(item.get("finding", "")).strip()
        if step in ADJUDICATION_EVIDENCE_STEPS and _is_concrete_finding(finding):
            return True, ""
    return (
        False,
        f"{decision} decision has no concrete evidence: every finding is empty, "
        "too short, or only references generic review language",
    )


def _format_scrub_findings(scrub_result) -> str:
    """Renders the coding scrubber's issue list as a paragraph the LLM can read."""
    if not isinstance(scrub_result, dict):
        return "No coding scrubber findings are available for this claim."
    issues = scrub_result.get("issues") or []
    lines = []
    for it in issues:
        if not isinstance(it, dict):
            continue
        sev = str(it.get("severity", "info")).upper()
        field = it.get("field", "-")
        msg = it.get("message", "")
        line = f"- [{sev}] {field}: {msg}"
        if it.get("suggestion"):
            line += f"  (suggested fix: {it['suggestion']})"
        lines.append(line)
    if lines:
        return "\n".join(lines)
    score = scrub_result.get("overall_score")
    if score is not None:
        return f"The coding scrubber flagged no issues (cleanliness score {score}/100)."
    return "No coding scrubber findings are available for this claim."


async def run_claim_adjudication_llm(claim: dict, insurer_id: str, fraud_result: dict) -> dict:
    """LLM adjudication for a low-fraud-risk claim whose payer has a policy on
    file. Returns {decision, rationale, evidence, policy_basis, retrieved_sources}
    where decision ∈ accept | deny | escalate.

    Never raises — on an unrecoverable LLM failure it returns a safe `escalate`
    verdict so the claim reaches a human rather than being auto-accepted/denied."""
    diagnosis_codes = claim.get("diagnosis_codes") or []
    procedure_codes = claim.get("procedure_codes") or []

    claim_view = {
        "patient_name": claim.get("patient_name"),
        "patient_id": claim.get("patient_id"),
        "date_of_service": claim.get("date_of_service"),
        "provider_name": claim.get("provider_name"),
        "payer_name": claim.get("payer_name"),
        "diagnosis_codes": diagnosis_codes,
        "procedure_codes": procedure_codes,
        "billed_amount": claim.get("total_billed"),
        "currency": claim.get("currency"),
        "pre_auth_number": claim.get("pre_auth_number"),
        "notes": claim.get("notes"),
    }

    auth_check_summary = _format_auth_check({
        "status": claim.get("auth_check_status"),
        "detail": claim.get("auth_check_detail"),
    })
    scrub_findings = _format_scrub_findings(claim.get("scrub_result"))

    # RAG: retrieve the payer's own policy rules for these codes.
    policy_context = "No specific payer rules were retrieved for these codes."
    retrieved_chunks = []
    try:
        search_query = (
            f"Coverage, exclusions and billing rules for diagnoses "
            f"{', '.join(diagnosis_codes)} and procedures {', '.join(procedure_codes)}"
        )
        embeddings_model = get_embeddings()
        query_vector = embeddings_model.embed_query(search_query)
        res = supabase.rpc("match_policy_rules", {
            "query_embedding": query_vector,
            "match_threshold": 0.4,
            "match_count": 5,
            "p_insurer_id": insurer_id,
        }).execute()
        if res.data:
            retrieved_chunks = [m["content"] for m in res.data]
            policy_context = "\n---\n".join(retrieved_chunks)
    except Exception as e:
        logger.error(f"Adjudication RAG retrieval failed: {e}")

    system_prompt = ADJUDICATION_SYSTEM_PROMPT.format(
        auth_check_summary=auth_check_summary,
        scrub_findings=scrub_findings,
        policy_context=policy_context,
    )
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=(
            "Adjudicate the following claim and return the final decision:\n\n"
            + json.dumps(claim_view, indent=2, default=str)
        )),
    ]
    llm = get_llm(json_mode=True)

    try:
        response = await llm.ainvoke(messages)
        result = parse_llm_json(response.content)
    except Exception as e:
        logger.error(f"Adjudication LLM call/parse failed: {e}")
        response, result = None, {}

    is_valid, reason = _validate_adjudication(result)
    if not is_valid and response is not None:
        logger.warning(f"Adjudication decision rejected ({reason}). Re-prompting once.")
        try:
            correction = HumanMessage(content=(
                f"Your previous response was rejected: {reason}.\n\n"
                "Return ONLY corrected JSON. `decision` must be accept, deny, or "
                "escalate. A `deny` or `escalate` decision MUST include an "
                "`evidence` array whose `finding` names a SPECIFIC code, field, "
                "policy rule, or authorization verdict. Vague phrases like "
                "'further review' are NOT evidence. If you cannot cite something "
                "specific, choose accept."
            ))
            messages.extend([response, correction])
            response = await llm.ainvoke(messages)
            result = parse_llm_json(response.content)
            is_valid, reason = _validate_adjudication(result)
        except Exception as e:
            logger.error(f"Adjudication re-prompt failed: {e}")
            is_valid = False

    if not is_valid:
        # A claim we cannot adjudicate cleanly goes to a human — never silently
        # auto-accept or auto-deny on a parsing failure.
        logger.error(f"Adjudication still invalid ({reason}); falling back to escalate.")
        return {
            "decision": "escalate",
            "rationale": (
                "This claim has been routed to a human reviewer. The automated "
                "adjudication could not reach a reliable decision, so a reviewer "
                "should complete the assessment."
            ),
            "evidence": [],
            "policy_basis": [],
            "retrieved_sources": retrieved_chunks,
        }

    evidence = result.get("evidence")
    policy_basis = result.get("policy_basis")
    return {
        "decision": str(result.get("decision", "")).strip().lower(),
        "rationale": result.get("rationale") or "No rationale provided.",
        "evidence": evidence if isinstance(evidence, list) else [],
        "policy_basis": policy_basis if isinstance(policy_basis, list) else [],
        "retrieved_sources": retrieved_chunks,
    }


async def generate_medical_recommendation(claim_data: dict) -> dict:
    """Produces a structured, ADVISORY medical-necessity opinion for a claim.

    This never decides accept/deny — claim adjudication (services/adjudication.py)
    and the human medical officer own that. It assembles the real clinical
    context for the claim (code descriptions, provider notes, any linked
    pre-auth document text, payer-policy RAG, and the scrub/auth/fraud signals
    already computed) and returns the structured verdict persisted to
    `claims.medical_necessity`.
    """
    claim_id = claim_data.get("id")
    logger.info(f"Generating medical-necessity review for claim {claim_id}")

    # --- 1. Clinical codes, with catalogue descriptions ----------------------
    raw_dx = claim_data.get("diagnosis_codes") or []
    raw_px = claim_data.get("procedure_codes") or []
    if isinstance(raw_dx, str):
        raw_dx = [raw_dx]
    if isinstance(raw_px, str):
        raw_px = [raw_px]
    dx_lines = [format_code_with_description(c, describe_diagnosis(c)) for c in raw_dx if c]
    px_lines = [format_code_with_description(c, describe_procedure(c)) for c in raw_px if c]
    dx_block = "\n".join(f"- {ln}" for ln in dx_lines) or "- None supplied"
    px_block = "\n".join(f"- {ln}" for ln in px_lines) or "- None supplied"

    # --- 2. Provider clinical context ----------------------------------------
    notes = (claim_data.get("notes") or "").strip()
    amount = claim_data.get("total_billed") or claim_data.get("billed_amount") or 0

    # --- 3. Clinical documents from a linked pre-authorisation ---------------
    doc_text = ""
    pre_auth_id = claim_data.get("pre_auth_id")
    if pre_auth_id:
        try:
            docs_res = supabase.table("pre_auth_documents").select(
                "file_name, extracted_text"
            ).eq("pre_auth_id", pre_auth_id).execute()
            for d in (docs_res.data or []):
                txt = (d.get("extracted_text") or "").strip()
                if txt:
                    doc_text += f"--- {d.get('file_name', 'document')} ---\n{txt}\n\n"
        except Exception as e:
            logger.warning(f"Claim {claim_id}: could not load linked pre-auth documents: {e}")

    # --- 4. Signals already computed by the pipeline -------------------------
    scrub_issues = []
    scrub_result = claim_data.get("scrub_result")
    if isinstance(scrub_result, dict):
        for issue in (scrub_result.get("issues") or []):
            if isinstance(issue, dict):
                msg = issue.get("message") or issue.get("field") or ""
                if msg:
                    scrub_issues.append(f"[{issue.get('severity', 'info')}] {msg}")
    scrub_block = "\n".join(f"- {s}" for s in scrub_issues) or "- No coding issues recorded."

    auth_status = claim_data.get("auth_check_status") or "not_applicable"
    auth_detail = claim_data.get("auth_check_detail") or "No pre-authorisation referenced."

    fraud_level = claim_data.get("fraud_risk_level")
    fraud_flags = claim_data.get("fraud_flags") or []
    if isinstance(fraud_flags, list):
        fraud_flags = ", ".join(str(f) for f in fraud_flags if f)
    fraud_block = f"Risk level: {fraud_level or 'not yet screened'}" + (
        f"; flags: {fraud_flags}" if fraud_flags else ""
    )

    # --- 5. Payer policy RAG -------------------------------------------------
    policy_excerpts = []
    payer_id = claim_data.get("payer_id")
    if payer_id:
        try:
            query_text = "\n".join(
                filter(None, [", ".join(px_lines), ", ".join(dx_lines), notes[:600]])
            ) or "medical necessity coverage criteria"
            embeddings_model = get_embeddings()
            query_vector = embeddings_model.embed_query(query_text[:2000])
            policy_res = supabase.rpc("match_policy_rules", {
                "query_embedding": query_vector,
                "match_threshold": 0.3,
                "match_count": 5,
                "p_insurer_id": payer_id,
            }).execute()
            policy_excerpts = [m["content"] for m in (policy_res.data or []) if m.get("content")]
        except Exception as e:
            logger.warning(f"Claim {claim_id}: policy RAG retrieval failed: {e}")
    if policy_excerpts:
        policy_block = "\n\n".join(
            f"[Policy excerpt {i + 1}]\n{p}" for i, p in enumerate(policy_excerpts)
        )
    else:
        policy_block = (
            "No payer policy is on file for this insurer. Judge medical necessity "
            "against standard MCG/InterQual-style and WHO/NICE guidance, and state "
            "in your reasoning that no payer-specific policy was available."
        )

    # --- 6. Build the prompt -------------------------------------------------
    human_context = f"""## CLAIM UNDER REVIEW
- Patient: {claim_data.get('patient_name', 'Unknown')} (ID: {claim_data.get('patient_id', 'Unknown')})
- Date of service: {claim_data.get('date_of_service', 'Unknown')}
- Provider: {claim_data.get('provider_name', 'Unknown')}
- Billed amount: {amount}

### Billed diagnosis codes (ICD-10)
{dx_block}

### Billed procedure codes (CPT)
{px_block}

### Provider clinical notes
{notes or "(none provided)"}

### Clinical documents (from linked pre-authorisation)
{doc_text.strip() or "(no clinical documents attached)"}

## SYSTEM SIGNALS (context only — not your verdict)
### Coding-scrubber issues
{scrub_block}

### Authorisation linkage check
Status: {auth_status} — {auth_detail}

### Fraud screening
{fraud_block}

## PAYER POLICY EXCERPTS
{policy_block}
"""

    messages = [
        SystemMessage(content=MEDICAL_NECESSITY_PROMPT),
        HumanMessage(content=human_context),
    ]
    llm = get_llm(json_mode=True)

    try:
        response = await llm.ainvoke(messages)
        result = _safe_parse_json(response.content)
    except Exception as e:
        logger.error(f"Claim {claim_id}: medical-necessity LLM call failed: {e}")
        return _medical_necessity_fallback()

    is_valid, reason = _validate_medical_necessity(result)
    if not is_valid:
        logger.warning(
            f"Claim {claim_id}: medical-necessity result rejected ({reason}); re-prompting once."
        )
        correction = HumanMessage(content=(
            f"Your previous response was rejected: {reason}.\n\n"
            "Return ONLY corrected JSON in the exact required shape. `assessment` "
            "must be one of supported, partially_supported, "
            "insufficient_documentation, not_supported. `criteria` must be a "
            "non-empty array, and at least one entry must have a valid status and "
            "a concrete, specific `finding`."
        ))
        try:
            messages.extend([response, correction])
            response = await llm.ainvoke(messages)
            result = _safe_parse_json(response.content)
            is_valid, reason = _validate_medical_necessity(result)
        except Exception as e:
            logger.error(f"Claim {claim_id}: medical-necessity re-prompt failed: {e}")
            return _medical_necessity_fallback()

    if not is_valid:
        logger.error(
            f"Claim {claim_id}: medical-necessity still invalid after re-prompt ({reason})."
        )
        return _medical_necessity_fallback()

    # Normalise the validated LLM output to the stored shape.
    return {
        "assessment": str(result.get("assessment", "")).strip().lower(),
        "headline": str(result.get("headline", "")).strip(),
        "summary": str(result.get("summary", "")).strip(),
        "criteria": [
            {
                "name": str(c.get("name", "")).strip(),
                "status": str(c.get("status", "unknown")).strip().lower(),
                "finding": str(c.get("finding", "")).strip(),
            }
            for c in result.get("criteria", [])
            if isinstance(c, dict) and str(c.get("finding", "")).strip()
        ],
        "policy_basis": [
            {
                "snippet": str(p.get("snippet", "")).strip(),
                "note": str(p.get("note", "")).strip(),
            }
            for p in (result.get("policy_basis") or [])
            if isinstance(p, dict) and str(p.get("snippet", "")).strip()
        ],
        "recommended_actions": [
            str(a).strip() for a in (result.get("recommended_actions") or []) if str(a).strip()
        ],
        "policy_available": bool(policy_excerpts),
    }